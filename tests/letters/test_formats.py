"""Tests for the cover-letter PDF and DOCX converters.

The integration tests in test_letters.py cover the route surface (status,
content-type, magic header). These unit tests cover the rendering logic
that's most likely to break: the multi-line contact header with bold
name + plain contact line + markdown link line.
"""

from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document

from role_tracker.letters.formats import letter_to_docx, letter_to_pdf

_HEADER_LETTER = (
    "**Shaikh Mushfikur Rahman**\n"
    "+1 555 0100 | shaikh@example.com | Halifax, Canada\n"
    "https://linkedin.com/in/shaikh | "
    "https://github.com/shaikh | "
    "https://shaikh.dev"
    "\n\n"
    "Dear Acme team,\n\n"
    "I'm writing because **the McKesson supply chain project** taught me "
    "lessons that map to your platform.\n\n"
    "Best,\n"
    "Shaikh"
)


def test_pdf_renders_and_returns_bytes() -> None:
    pdf_bytes = letter_to_pdf(_HEADER_LETTER)
    assert isinstance(pdf_bytes, bytes)
    assert pdf_bytes.startswith(b"%PDF-")
    assert len(pdf_bytes) > 1000  # rough sanity — letter should be > 1KB


def test_pdf_with_page_count_returns_tuple() -> None:
    pdf_bytes, page_count = letter_to_pdf(
        _HEADER_LETTER, with_page_count=True
    )
    assert pdf_bytes.startswith(b"%PDF-")
    assert page_count == 1  # the sample letter fits on one page


def test_pdf_overflows_to_two_pages_with_long_text() -> None:
    """A letter long enough to spill should report page_count >= 2 so
    the route can surface a warning to the user."""
    long_body = (
        "**Jane Doe**\n"
        "+1 555 0100 | jane@example.com\n\n"
        "Dear Acme Team,\n\n"
        + ("Lorem ipsum dolor sit amet. " * 200 + "\n\n") * 4
        + "Best,\nJane"
    )
    _, page_count = letter_to_pdf(long_body, with_page_count=True)
    assert page_count >= 2


def test_pdf_renders_urls_as_plain_text() -> None:
    """ATS-friendly: URLs render as plain text in the PDF, not as
    markdown link syntax. The bare URL is what scrapers extract."""
    pdf_bytes = letter_to_pdf(_HEADER_LETTER)
    text = _extract_text_from_pdf_bytes(pdf_bytes)
    assert "https://linkedin.com/in/shaikh" in text
    assert "https://github.com/shaikh" in text
    # No leftover markdown link syntax (e.g. from older saved letters).
    assert "](" not in text


def test_docx_renders_urls_as_plain_text() -> None:
    docx_bytes = letter_to_docx(_HEADER_LETTER)
    full_text = _extract_text_from_docx_bytes(docx_bytes)
    assert "https://linkedin.com/in/shaikh" in full_text
    assert "](" not in full_text


def test_docx_preserves_multi_line_header_inside_one_paragraph() -> None:
    """The three header lines should share one paragraph with soft breaks
    between them — not three separate paragraphs (which would add extra
    vertical space) and not one collapsed line (which would look broken)."""
    docx_bytes = letter_to_docx(_HEADER_LETTER)
    doc = Document(io.BytesIO(docx_bytes))
    # First paragraph is the contact header. Each <w:br/> inside a paragraph
    # is a soft line break; we expect 2 of them (between 3 lines).
    first = doc.paragraphs[0]
    breaks = first._p.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
    )
    assert len(breaks) == 2


def test_docx_renders_name_as_bold() -> None:
    docx_bytes = letter_to_docx(_HEADER_LETTER)
    doc = Document(io.BytesIO(docx_bytes))
    first = doc.paragraphs[0]
    # The very first run should be the bold name.
    assert first.runs[0].text == "Shaikh Mushfikur Rahman"
    assert first.runs[0].bold is True


def test_docx_inline_bold_inside_body_paragraph() -> None:
    """`**bold**` mid-paragraph still becomes a bold run."""
    docx_bytes = letter_to_docx(_HEADER_LETTER)
    doc = Document(io.BytesIO(docx_bytes))
    # Find the body paragraph that has the bolded project name.
    body = next(
        p for p in doc.paragraphs if "McKesson" in p.text
    )
    bold_runs = [r for r in body.runs if r.bold]
    assert any("McKesson" in r.text for r in bold_runs)


def test_empty_blank_lines_dont_create_phantom_paragraphs() -> None:
    """Multiple blank lines collapse to a single paragraph break."""
    text = "**Name**\n\n\n\nDear team,\n\nBest,"
    docx_bytes = letter_to_docx(text)
    doc = Document(io.BytesIO(docx_bytes))
    # Three paragraphs: header, body, signoff. Not four or five.
    assert len(doc.paragraphs) == 3


# ----- helpers -----


def _extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Pull all visible text out of a PDF using pypdf (already a dep)."""
    pypdf = pytest.importorskip("pypdf")
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Read every <w:t> element in the document.xml of the .docx zip."""
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return xml
