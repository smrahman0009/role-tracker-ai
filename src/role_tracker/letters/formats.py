"""Cover-letter format converters.

The agent stores letters as Markdown-flavoured plain text — paragraphs
separated by blank lines, with `**bold**` markers around the contact
header (name + email + phone). For employer apply forms we need PDF or
DOCX, since `.md` is essentially never accepted.

Two converters here, both pure-Python with no system deps:
- letter_to_pdf: uses fpdf2 (Helvetica, 11pt, 1in margins, US Letter).
- letter_to_docx: uses python-docx (Calibri, 11pt, 1in margins).

Both render `**bold**` as actual bold runs. Other Markdown is passed
through as literal text — cover letters don't typically use lists,
tables, or other rich formatting.
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF


_BOLD_SPLIT = re.compile(r"(\*\*[^\n*]+?\*\*)")
# Markdown link `[label](url)` — we keep only `label` since neither fpdf2's
# basic API nor python-docx's `add_run` make hyperlinks easy. The URLs
# already live elsewhere on the page (Apply Kit profile fields).
_MD_LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def _strip_md_links(text: str) -> str:
    return _MD_LINK.sub(r"\1", text)


def _split_bold(line: str) -> list[tuple[str, bool]]:
    """Split a single line into (text, is_bold) chunks.

    Markdown `**X**` becomes a bold chunk; everything else stays plain.
    Empty chunks are filtered out. Caller is responsible for splitting
    multi-line input on `\n` first — this function does not handle line
    breaks itself.
    """
    parts: list[tuple[str, bool]] = []
    for chunk in _BOLD_SPLIT.split(line):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
            parts.append((chunk[2:-2], True))
        else:
            parts.append((chunk, False))
    return parts


def _normalize(text: str) -> list[list[str]]:
    """Turn the letter text into a list of paragraphs, each a list of lines.

    Letters use blank lines between paragraphs, and single newlines inside
    the contact header to separate `**Name**`, the contact line, and the
    links line. We need to honour both.
    """
    raw_paragraphs = [p for p in text.split("\n\n") if p.strip()]
    paragraphs: list[list[str]] = []
    for para in raw_paragraphs:
        para = _strip_md_links(para)
        lines = [ln for ln in para.split("\n") if ln.strip()]
        if lines:
            paragraphs.append(lines)
    return paragraphs


# ----- PDF -----


def letter_to_pdf(
    text: str, *, with_page_count: bool = False
) -> bytes | tuple[bytes, int]:
    """Render the letter text as a US-Letter PDF.

    Returns the PDF bytes by default. If `with_page_count=True`, returns
    `(bytes, pages)` so callers can detect overflow and warn the user.
    Auto-page-break is enabled, so a long letter will silently span two
    pages — the page count is the only signal the caller has.
    """
    pdf = FPDF(format="Letter", unit="pt")
    pdf.set_margins(left=72, top=72, right=72)  # 1 inch margins
    pdf.set_auto_page_break(auto=True, margin=72)
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)

    paragraphs = _normalize(text)
    line_height = 14  # ~1.27 leading at 11pt
    paragraph_gap = 8

    for i, lines in enumerate(paragraphs):
        for j, line in enumerate(lines):
            for run_text, is_bold in _split_bold(line):
                pdf.set_font("Helvetica", style="B" if is_bold else "", size=11)
                pdf.write(line_height, run_text)
            # End of line — drop to the next baseline. This handles both
            # the single-line breaks inside the header (Name / contacts /
            # links) and the end-of-paragraph break for body paragraphs.
            pdf.ln(line_height)
        if i < len(paragraphs) - 1:
            pdf.ln(paragraph_gap)

    output = bytes(pdf.output())
    if with_page_count:
        return output, pdf.page_no()
    return output


# ----- DOCX -----


def letter_to_docx(text: str) -> bytes:
    """Render the letter text as a .docx and return the bytes."""
    doc = Document()
    # 1 inch margins on all sides.
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    # Default style: Calibri 11pt (Word's standard).
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for lines in _normalize(text):
        p = doc.add_paragraph()
        for j, line in enumerate(lines):
            if j > 0:
                # Soft line break inside the same paragraph — preserves
                # the multi-line contact header without inserting the
                # extra spacing of a fresh <w:p>.
                p.add_run().add_break()
            for run_text, is_bold in _split_bold(line):
                run = p.add_run(run_text)
                run.bold = is_bold

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
